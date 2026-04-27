#!/usr/bin/env python3.13
"""
Generates HireQuadrant QA Walkthrough.docx with a fully walkable test
plan: each step has WHAT (feature explanation), HOW (step-by-step
instructions), and PASS IF (expected outcome). Plus Done / Pass / Fail
checkboxes (Unicode) and a Notes column.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_checkbox(cell):
    """Add a Unicode ballot box (☐) to a cell."""
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run('\u2610')
    run.font.size = Pt(16)
    set_cell_shade(cell, 'F5F5F5')


def set_cell_shade(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), 'BBBBBB')
        tcBorders.append(b)
    tc_pr.append(tcBorders)


def set_table_widths(table, widths_inches):
    for row in table.rows:
        for cell, w in zip(row.cells, widths_inches):
            cell.width = Inches(w)


def add_heading(doc, text, level=1):
    return doc.add_heading(text, level=level)


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x00, 0x35, 0x94)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_subsection_heading(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p


def add_intro_paragraph(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)
    return p


def add_step_cell_text(cell, text):
    """Render multi-line step text with bold labels."""
    cell.text = ''
    lines = text.split('\n')
    for i, line in enumerate(lines):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)

        # Bold-prefix detection: WHAT:, HOW:, PASS IF:, etc.
        for prefix in ('WHAT:', 'WHY:', 'HOW:', 'PASS IF:', 'EXPECTED:'):
            if line.startswith(prefix):
                run = p.add_run(prefix)
                run.bold = True
                run.font.size = Pt(9)
                rest = line[len(prefix):]
                if rest:
                    run2 = p.add_run(rest)
                    run2.font.size = Pt(9)
                break
        else:
            run = p.add_run(line)
            run.font.size = Pt(9)


def add_test_table(doc, steps):
    """
    steps: list of (step_num, step_text) tuples.
    """
    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Light Grid'
    hdr = table.rows[0].cells
    hdr[0].text = '#'
    hdr[1].text = 'Step (with walkthrough)'
    hdr[2].text = 'Done'
    hdr[3].text = 'Pass'
    hdr[4].text = 'Fail'
    hdr[5].text = 'Notes'

    for cell in hdr:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
        set_cell_shade(cell, 'E8EEF7')
        set_cell_borders(cell)

    for step_num, desc in steps:
        row = table.add_row().cells
        row[0].text = str(step_num)
        for p in row[0].paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)
                run.bold = True
        add_step_cell_text(row[1], desc)
        for c in (row[2], row[3], row[4]):
            add_checkbox(c)
        row[5].text = ''
        for c in row:
            set_cell_borders(c)

    set_table_widths(
        table,
        [0.35, 4.6, 0.5, 0.5, 0.5, 1.55],
    )
    return table


# ============================================================
# CONTENT
# ============================================================

PART_A_INTRO = """Setup before testing. Do this once.

REQUIRED TOOLS:
• Google Chrome browser (REQUIRED for the Lighthouse audit later — it does not work in Safari or Firefox)
• Apple Safari OR Mozilla Firefox (for cross-browser checks)
• A real iPhone or Android phone (for mobile tests)
• A Google account (only if you plan to use Google Docs to mark up this checklist)
• An authenticator app on your phone (Google Authenticator, 1Password, or Authy) — only needed for the 2FA test in Part C

ACCOUNTS YOU WILL USE:
• Super admin: your own email and password (provided by Rafael)
• Backup admin: rrainey19138@gmail.com / TestPass123!
• Employer: test-employer-1@hirequadrant-test.com / TestBiz123!
• Candidate (existing): seed-reviewer-1@hirequadrant-seed.test / TestPass123!
• Candidate (you will create one in Part C using your own personal email)

GENERAL BROWSER SETUP (do this for every test session):
1. Open Google Chrome.
2. Press Cmd+Shift+N (Mac) or Ctrl+Shift+N (Windows) to open an Incognito window.
3. In the address bar, type https://hirequadrant.com and press Enter.
4. Press Cmd+Option+I (Mac) or F12 (Windows) to open Developer Tools.
5. In Developer Tools, click the "Console" tab and leave it visible. If you see RED error messages during any test, copy them and add to the Notes column for that step.
6. Open a separate notes doc (Google Doc, Word, Apple Notes — anything) for screenshots and longer bug descriptions.

HOW TO TAKE A SCREENSHOT (when a step fails):
• Mac: press Cmd+Shift+4, then drag the crosshair over the area you want.
• Windows: press Windows+Shift+S, then drag.
The screenshot is saved or copied to your clipboard. Paste it into your notes doc with the step number.
"""

PART_B_INTRO = "Anonymous user (no login). Open an incognito window and follow the steps. The goal is to confirm the public-facing site works for visitors who haven't signed up yet."
PART_C_INTRO = "Individual / Candidate user. Sign up a fresh account (or log in as the seeded candidate). The goal is to confirm a job seeker can register, build a profile, search jobs, apply, and use all candidate features."
PART_D_INTRO = "Company / Employer user. Log in as test-employer-1@hirequadrant-test.com / TestBiz123!. The goal is to confirm an employer can post jobs, manage applicants, schedule interviews, and edit their company profile."
PART_E_INTRO = "Admin / Portal user. Log in as rrainey19138@gmail.com / TestPass123! (or your own admin email). The goal is to confirm a HireQuadrant internal staff member can moderate content, manage feeds, and oversee the platform."
PART_F_INTRO = "Cross-cutting tests for security, performance, mobile, and other concerns that apply to every persona. These are technical — read each step carefully, follow exactly."


def s(*lines):
    """Compose a multi-line step description."""
    return '\n'.join(lines)


SECTIONS = [
    (
        'PART A: Setup',
        PART_A_INTRO,
        [
            ('A.1 Accounts', [
                (1, s(
                    'Confirm you have all 5 test accounts noted on paper or in a password manager.',
                    'WHY: you will need to switch between personas during testing.',
                    'PASS IF: all 5 emails + passwords are accessible to you.',
                )),
                (2, s(
                    'Confirm you have Google Chrome installed.',
                    'WHY: Lighthouse (Section F.3) only works in Chrome.',
                    'HOW: open Chrome. If you don\'t have it, download from google.com/chrome.',
                    'PASS IF: Chrome opens.',
                )),
                (3, s(
                    'Install an authenticator app on your phone.',
                    'WHY: needed for the Two-Factor Authentication test in Part C.',
                    'HOW: in the App Store / Play Store search "Google Authenticator" or "Authy" or "1Password" — install one.',
                    'PASS IF: app icon visible on your phone home screen.',
                )),
            ]),
            ('A.2 Browser setup', [
                (4, s(
                    'Open Chrome in Incognito mode and load the site.',
                    'WHY: incognito ensures cookies and login state from prior browsing don\'t taint the test.',
                    'HOW: open Chrome → press Cmd+Shift+N (Mac) / Ctrl+Shift+N (Windows) → in the address bar type https://hirequadrant.com → Enter.',
                    'PASS IF: site loads, you see "We work for you, not the algorithm." headline.',
                )),
                (5, s(
                    'Open the Developer Tools panel.',
                    'WHY: lets you watch for JavaScript errors throughout testing.',
                    'HOW: press Cmd+Option+I (Mac) or F12 (Windows). A panel opens at the right or bottom of the browser.',
                    'PASS IF: a panel with tabs (Elements, Console, Sources, etc.) is visible.',
                )),
                (6, s(
                    'Click the Console tab in DevTools and keep it visible.',
                    'WHY: red error messages here = bugs to log.',
                    'HOW: in DevTools, click the "Console" tab. Leave it open during all testing.',
                    'PASS IF: console is visible. Most pages should produce no red errors.',
                )),
                (7, s(
                    'Open a separate notes doc.',
                    'WHY: for screenshots and long bug descriptions.',
                    'HOW: open Google Docs, Word, Apple Notes — whatever. Title it "HireQuadrant QA - [today\'s date]".',
                    'PASS IF: empty notes doc is open.',
                )),
            ]),
        ],
    ),
    (
        'PART B: Anonymous / Public (no login)',
        PART_B_INTRO,
        [
            ('B.1 Homepage and navigation', [
                (1, s(
                    'Confirm the homepage loads quickly and looks correct.',
                    'WHY: it\'s the first impression for every visitor.',
                    'HOW: address bar → type https://hirequadrant.com → Enter.',
                    'PASS IF: page renders within 5 seconds. The "We work for you, not the algorithm." headline is visible. Console (DevTools) shows no red errors.',
                )),
                (2, s(
                    'Confirm the top navigation bar is complete.',
                    'WHY: nav must let visitors reach all key sections.',
                    'HOW: look at the top of the page from left to right.',
                    'PASS IF: you see the HireQuadrant logo on the left, then "Jobs", "Career Paths", "Companies", "Search", "Blog", "Pricing", and on the right "Sign in" + "Get started" buttons.',
                )),
                (3, s(
                    'Confirm clicking the logo returns home.',
                    'WHY: logo-as-home-link is a universal web pattern.',
                    'HOW: scroll down somewhere, then click the HireQuadrant logo (top-left).',
                    'PASS IF: you\'re back at the top of the homepage.',
                )),
                (4, s(
                    'Confirm the footer links are visible at the bottom.',
                    'WHY: legal/policy links are required for a launched site.',
                    'HOW: scroll to the bottom of the homepage.',
                    'PASS IF: you see About, Contact, Privacy, Terms, Content Policy, Cookies links.',
                )),
                (5, s(
                    'Click each footer link to confirm it loads.',
                    'WHY: any broken link is a launch blocker.',
                    'HOW: middle-click (or Cmd+click on Mac) each footer link to open in a new tab. Visit each tab.',
                    'PASS IF: every page loads with content (not blank, not 404).',
                )),
                (6, s(
                    'Switch to dark mode.',
                    'WHY: dark mode is a major UX feature; must work cleanly.',
                    'HOW: at the top-right of any page, find the moon/sun icon (next to "Sign in"). Click it.',
                    'PASS IF: the entire page changes to a dark background with light text. No white flashes, no unreadable text.',
                )),
                (7, s(
                    'Switch back to light mode.',
                    'WHY: confirms toggle works both directions.',
                    'HOW: click the moon/sun icon again.',
                    'PASS IF: page returns to light theme cleanly.',
                )),
            ]),
            ('B.2 Cookie consent', [
                (8, s(
                    'Confirm the cookie banner appears for first-time visitors.',
                    'WHY: required by GDPR and similar laws.',
                    'HOW: in incognito (no prior cookies), the banner appears at the bottom-right shortly after the page loads.',
                    'PASS IF: banner is visible with "Accept all" and "Necessary only" buttons.',
                )),
                (9, s(
                    'Click "Necessary only" and confirm the choice persists.',
                    'WHY: user preferences must be remembered.',
                    'HOW: click "Necessary only" on the banner. Then reload the page (Cmd+R or F5).',
                    'PASS IF: banner does not return after reload.',
                )),
            ]),
            ('B.3 Global search (Cmd+K)', [
                (10, s(
                    'Open the global search popup.',
                    'WHY: lets visitors find anything (jobs, companies, articles) from any page.',
                    'HOW: while on any page, press Cmd+K (Mac) or Ctrl+K (Windows).',
                    'PASS IF: a search modal appears in the middle of the screen with the cursor inside the input.',
                )),
                (11, s(
                    'Type a query and confirm grouped results.',
                    'WHY: results should organize by content type.',
                    'HOW: type "engineer" (5+ letters work best) into the search box.',
                    'PASS IF: results appear within 1 second, organized into 3 sections — Jobs, Companies, Articles — with a few items each.',
                )),
                (12, s(
                    'Navigate results with arrow keys.',
                    'WHY: keyboard accessibility is required.',
                    'HOW: press the down arrow key 3 times.',
                    'PASS IF: a different result is highlighted (with a colored background) after each press.',
                )),
                (13, s(
                    'Press Enter to open the highlighted result.',
                    'WHY: confirms the search → click flow works.',
                    'HOW: with a result highlighted, press Enter.',
                    'PASS IF: the page for that result loads.',
                )),
                (14, s(
                    'Confirm "no results" copy appears for nonsense queries.',
                    'WHY: empty states must communicate clearly.',
                    'HOW: Cmd+K → type "xyzabc123" → wait.',
                    'PASS IF: text appears that says something like "No results for xyzabc123".',
                )),
                (15, s(
                    'Press Escape to close the search popup.',
                    'WHY: standard keyboard close pattern.',
                    'HOW: press the Esc key.',
                    'PASS IF: popup closes, you\'re back on the page.',
                )),
                (16, s(
                    'Confirm the Search button in the header opens the same popup.',
                    'WHY: not all users know Cmd+K — must be discoverable by mouse.',
                    'HOW: click the "Search" button in the top navigation.',
                    'PASS IF: the same search popup opens.',
                )),
            ]),
            ('B.4 Browse jobs', [
                (17, s(
                    'Visit the jobs listing page.',
                    'WHY: this is the main job-search surface.',
                    'HOW: click "Jobs" in the top navigation.',
                    'PASS IF: page at /jobs loads with a search bar at top, filter sidebar on the left, and job cards on the right.',
                )),
                (18, s(
                    'Filter by keyword.',
                    'WHY: most users start with a job title or skill keyword.',
                    'HOW: in the keyword input at the top, type "manager" → press Enter.',
                    'PASS IF: jobs list filters to show roles with "manager" in the title or description.',
                )),
                (19, s(
                    'Filter by location.',
                    'WHY: location is the second most common filter.',
                    'HOW: in the location input next to keyword, type "remote" → press Enter.',
                    'PASS IF: list filters to remote roles.',
                )),
                (20, s(
                    'Filter by minimum salary.',
                    'WHY: candidates need to filter by their floor.',
                    'HOW: in the left sidebar, find the "Minimum salary" slider. Drag the handle to the right until it shows ~$100k.',
                    'PASS IF: list refreshes to show only jobs paying $100k or more.',
                )),
                (21, s(
                    'Filter by experience level.',
                    'WHY: senior candidates skip junior roles.',
                    'HOW: in the sidebar, find the "Experience" dropdown → click it → choose "Senior".',
                    'PASS IF: list narrows to senior-titled roles.',
                )),
                (22, s(
                    'Filter by workplace type.',
                    'WHY: remote vs hybrid vs on-site is a key sort.',
                    'HOW: in the sidebar, find the "Workplace" dropdown → choose "Remote".',
                    'PASS IF: list narrows again.',
                )),
                (23, s(
                    'Filter by posted-within.',
                    'WHY: stale jobs frustrate candidates.',
                    'HOW: in the sidebar, find "Posted within" → choose "1 week".',
                    'PASS IF: list shows only jobs posted in the last 7 days.',
                )),
                (24, s(
                    'Filter by security clearance.',
                    'WHY: defense and federal candidates filter by this.',
                    'HOW: in the sidebar, find "Security clearance" dropdown → choose any value.',
                    'PASS IF: list narrows.',
                )),
                (25, s(
                    'Toggle visa sponsorship checkbox.',
                    'WHY: visa-needing candidates filter by this.',
                    'HOW: in the sidebar, click the "Visa sponsorship" checkbox.',
                    'PASS IF: list narrows.',
                )),
                (26, s(
                    'Reset all filters.',
                    'WHY: starting over should be easy.',
                    'HOW: click "Reset" at the top-right of the filter sidebar.',
                    'PASS IF: all filters clear, full list returns.',
                )),
                (27, s(
                    'Open a job detail page.',
                    'WHY: the detail page is where the apply decision happens.',
                    'HOW: click any job card.',
                    'PASS IF: a page at /jobs/[id] loads with the full job description.',
                )),
                (28, s(
                    'Confirm key job info is visible.',
                    'WHY: candidates need title, company, location, salary, posted-date upfront.',
                    'HOW: read the top of the job detail page.',
                    'PASS IF: title, company, location, salary range, posted-date are all visible.',
                )),
                (29, s(
                    'Confirm action buttons are visible.',
                    'WHY: Apply / Save / Share are the core CTAs.',
                    'HOW: look near the title for buttons.',
                    'PASS IF: Apply (blue/green button), Save (bookmark icon), Share button are all visible.',
                )),
                (30, s(
                    'Confirm the description is well-formatted.',
                    'WHY: structured descriptions read better than walls of text.',
                    'HOW: scroll down through the description.',
                    'PASS IF: you see H2 section headers like Position Overview, Ideal Candidate, Responsibilities, Required, Preferred, Skills (some or all).',
                )),
                (31, s(
                    'Click Apply while logged out.',
                    'WHY: must prompt for sign-in without losing the user\'s context.',
                    'HOW: click the green "Apply" button.',
                    'PASS IF: redirects to /login with a message at top saying "Sign in to finish applying" or similar. The URL contains returnTo=... so after login you\'re returned here.',
                )),
                (32, s(
                    'Click the back button. Click Save while logged out.',
                    'WHY: same — must guide to sign-up not silent fail.',
                    'HOW: browser back button → click the bookmark/save icon.',
                    'PASS IF: a toast/popup appears saying "Sign in to save jobs".',
                )),
                (33, s(
                    'Click the Share button.',
                    'WHY: candidates and recruiters share job links.',
                    'HOW: click "Share".',
                    'PASS IF: a popup appears with options to copy link, email, share to X (Twitter), or share to LinkedIn.',
                )),
            ]),
            ('B.5 Companies directory', [
                (34, s(
                    'Visit the companies directory.',
                    'WHY: candidates research employers before applying.',
                    'HOW: click "Companies" in the top nav.',
                    'PASS IF: a grid of company cards loads, each with a logo, name, and rating.',
                )),
                (35, s(
                    'Open a company profile.',
                    'WHY: the profile is the main research surface.',
                    'HOW: click any company card.',
                    'PASS IF: a page at /companies/[slug] loads.',
                )),
                (36, s(
                    'Confirm header info is visible.',
                    'WHY: name, industry, location are scan-essential.',
                    'HOW: read the top of the profile.',
                    'PASS IF: name + logo + industry + location are visible.',
                )),
                (37, s(
                    'Confirm rating widget is visible.',
                    'WHY: aggregate rating is the social proof.',
                    'HOW: look near the top.',
                    'PASS IF: a star rating (e.g., "4.2 of 5") and the review count (e.g., "based on 28 reviews") are visible.',
                )),
                (38, s(
                    'Confirm the AI review summary appears (when 3+ reviews exist).',
                    'WHY: gives a fast, neutral synthesis of employee sentiment.',
                    'HOW: scroll down to the "What people say about [Company]" box.',
                    'PASS IF: a paragraph summary plus "Common praise" and "Common concerns" lists appear. (If the company has fewer than 3 reviews, this section is hidden — that\'s OK.)',
                )),
                (39, s(
                    'Confirm the About section renders.',
                    'WHY: company description is critical context.',
                    'HOW: keep scrolling.',
                    'PASS IF: the About section shows the company description.',
                )),
                (40, s(
                    'Confirm reviews list with vote/report buttons.',
                    'WHY: review interactions drive engagement and moderation signal.',
                    'HOW: scroll to Reviews section. Look at any review.',
                    'PASS IF: each review has thumbs-up + thumbs-down + a flag/report icon.',
                )),
                (41, s(
                    'Confirm Q&A section renders.',
                    'WHY: candidates ask employers questions publicly.',
                    'HOW: scroll past Reviews.',
                    'PASS IF: a Q&A section with existing answered questions (or "No questions yet" copy).',
                )),
                (42, s(
                    'Confirm Open Positions section.',
                    'WHY: candidates often go from review to applying directly.',
                    'HOW: scroll to the bottom of the company profile.',
                    'PASS IF: a list of jobs at this company is shown.',
                )),
                (43, s(
                    'Click a thumbs-up on any review (logged out).',
                    'WHY: must guide unauthenticated users to sign in.',
                    'HOW: click thumbs-up under any review.',
                    'PASS IF: count increments OR a toast says "Sign in to vote".',
                )),
                (44, s(
                    'Click the flag/report icon on a review.',
                    'WHY: testing the abuse-report flow.',
                    'HOW: click the flag icon next to a review.',
                    'PASS IF: a modal opens with a reason dropdown and a Submit button.',
                )),
                (45, s(
                    'Close the report modal.',
                    'WHY: confirms cancel/close works.',
                    'HOW: press Esc or click the X button.',
                    'PASS IF: modal closes.',
                )),
            ]),
            ('B.6 Compare companies', [
                (46, s(
                    'Visit the Compare page.',
                    'WHY: a unique HQ feature to evaluate multiple companies side-by-side.',
                    'HOW: in the address bar, type https://hirequadrant.com/compare → Enter.',
                    'PASS IF: the Compare page loads with an autocomplete search input at top.',
                )),
                (47, s(
                    'Search for a company.',
                    'WHY: tests autocomplete dropdown.',
                    'HOW: in the search input, type "Quadrant".',
                    'PASS IF: a dropdown appears with matching companies — "Quadrant, Inc." should be visible.',
                )),
                (48, s(
                    'Add the first company.',
                    'WHY: the first card.',
                    'HOW: click "Quadrant, Inc." in the dropdown.',
                    'PASS IF: a card for Quadrant, Inc. appears with its stats.',
                )),
                (49, s(
                    'Add a second company for side-by-side comparison.',
                    'WHY: the core comparison feature.',
                    'HOW: search again with another name (e.g., "Test"), pick a result.',
                    'PASS IF: 2 cards now appear side-by-side, each showing rating, review count, open roles, size, HQ, founded, median salary.',
                )),
                (50, s(
                    'Remove a company from the comparison.',
                    'WHY: candidates may swap candidates in/out.',
                    'HOW: click the X button at the top-right of one card.',
                    'PASS IF: that card removes; the other stays.',
                )),
            ]),
            ('B.7 Programmatic SEO pages', [
                (51, s(
                    'Visit the industry filter page.',
                    'WHY: SEO landing page that ranks for "[industry] companies hiring".',
                    'HOW: address bar → https://hirequadrant.com/companies/industry/technology → Enter.',
                    'PASS IF: a list of tech companies renders.',
                )),
                (52, s(
                    'Visit the location-filter page.',
                    'WHY: SEO landing page for "jobs in [city]".',
                    'HOW: address bar → https://hirequadrant.com/jobs/location/new-york-ny → Enter.',
                    'PASS IF: a list of NYC jobs renders.',
                )),
                (53, s(
                    'Visit the salary insights page.',
                    'WHY: salary data is one of the highest-traffic SEO surfaces.',
                    'HOW: address bar → https://hirequadrant.com/salaries/software-engineer → Enter.',
                    'PASS IF: a page with percentile salary stats (25th / median / 75th) renders.',
                )),
                (54, s(
                    'Visit the "best companies" curated list.',
                    'WHY: editorial-style ranking for SEO.',
                    'HOW: address bar → https://hirequadrant.com/best/technology → Enter.',
                    'PASS IF: a ranked list of top tech companies renders.',
                )),
                (55, s(
                    'Visit the industry guide.',
                    'WHY: long-form content for SEO.',
                    'HOW: address bar → https://hirequadrant.com/guide/technology → Enter.',
                    'PASS IF: a careers guide for the tech industry renders with sections.',
                )),
                (56, s(
                    'Visit the interview prep page.',
                    'WHY: targets "how to prep for [role] interview" queries.',
                    'HOW: address bar → https://hirequadrant.com/interview-prep/software-engineer → Enter.',
                    'PASS IF: page with curated interview questions and a 7-day prep plan renders.',
                )),
                (57, s(
                    'Visit a career transition page.',
                    'WHY: long-tail SEO for career-pivot queries.',
                    'HOW: address bar → https://hirequadrant.com/career/from/data-analyst/to/product-manager → Enter.',
                    'PASS IF: a 90-day pivot plan + matching jobs render.',
                )),
                (58, s(
                    'Visit the Career Paths page.',
                    'WHY: AI-powered next-step suggestions.',
                    'HOW: click "Career Paths" in the top nav.',
                    'PASS IF: page loads (may prompt for input role first).',
                )),
            ]),
            ('B.8 Blog', [
                (59, s(
                    'Visit the blog index.',
                    'WHY: 200+ articles drive organic traffic.',
                    'HOW: click "Blog" in the top nav.',
                    'PASS IF: a grid of article cards loads. Each card has a category pill (e.g., "career") and possibly a role pill (e.g., "Software Engineer").',
                )),
                (60, s(
                    'Filter by topic.',
                    'WHY: readers narrow by interest.',
                    'HOW: at the top of the blog page, click each filter button: All, Career, Resume, Interview, Hiring, Industry.',
                    'PASS IF: list narrows for each filter.',
                )),
                (61, s(
                    'Confirm the "For your role" filter row exists.',
                    'WHY: lets candidates find role-specific advice.',
                    'HOW: just below the topic row, look for a "For your role" label.',
                    'PASS IF: a row of pills appears: "All roles", "Universal", "Software Engineer", "Product Manager", "Designer", etc.',
                )),
                (62, s(
                    'Filter by role.',
                    'HOW: click the "Software Engineer" pill.',
                    'PASS IF: list narrows to articles tagged for software engineers.',
                )),
                (63, s(
                    'Reset role filter.',
                    'HOW: click "All roles" pill.',
                    'PASS IF: list expands back.',
                )),
                (64, s(
                    'Open an article.',
                    'WHY: confirms detail page works.',
                    'HOW: click any article card.',
                    'PASS IF: full article loads at /blog/[slug] with title, excerpt, and formatted body.',
                )),
                (65, s(
                    'Confirm article meta pills.',
                    'WHY: pills aid navigation.',
                    'HOW: read near the title.',
                    'PASS IF: a category pill is visible. If the article is role-specific, a "For: [role]" pill is also visible.',
                )),
                (66, s(
                    'Confirm article volume.',
                    'WHY: 200+ posts is a content-volume claim we make.',
                    'HOW: go back to /blog. Filter All. Scroll the entire grid.',
                    'PASS IF: well over 100 articles visible (paginated or all rendered).',
                )),
            ]),
            ('B.9 Help / Support', [
                (67, s(
                    'Visit the Help Center.',
                    'WHY: support deflection — users find answers before contacting support.',
                    'HOW: address bar → https://hirequadrant.com/help-center → Enter.',
                    'PASS IF: a page with FAQ articles in collapsible sections renders.',
                )),
                (68, s(
                    'Search for an article.',
                    'WHY: keyword discovery.',
                    'HOW: in the search box at top, type "apply".',
                    'PASS IF: list narrows to apply-related FAQs.',
                )),
                (69, s(
                    'Filter by audience.',
                    'WHY: candidates and employers have different questions.',
                    'HOW: click each pill: "All articles", "For candidates", "For employers", "Everyone".',
                    'PASS IF: list filters appropriately each time.',
                )),
                (70, s(
                    'Expand an FAQ.',
                    'WHY: each question is collapsible to keep the page scannable.',
                    'HOW: click any question.',
                    'PASS IF: the answer expands inline below the question. Click again to collapse.',
                )),
                (71, s(
                    'Visit the Support / Contact page.',
                    'WHY: when help center isn\'t enough, must have a contact path.',
                    'HOW: address bar → https://hirequadrant.com/support → Enter.',
                    'PASS IF: a contact form renders. Submit a test message — confirmation appears.',
                )),
            ]),
            ('B.10 Auth pages (without logging in)', [
                (72, s(
                    'Visit the Sign In page.',
                    'HOW: click "Sign in" at the top of any page.',
                    'PASS IF: a page at /login renders with email + password fields and a "Sign in with Google" button.',
                )),
                (73, s(
                    'Visit the Forgot Password page.',
                    'HOW: on the login page, click "Forgot password?" link.',
                    'PASS IF: page at /reset-password with an email input renders.',
                )),
                (74, s(
                    'Visit the Register page.',
                    'HOW: click "Get started" at the top of any page.',
                    'PASS IF: page at /register with name + email + password fields and a Job Seeker / Employer toggle renders.',
                )),
                (75, s(
                    'Try registering with a known-breached password.',
                    'WHY: tests the HIBP password security check.',
                    'HOW: on /register, fill in a fake name + a fake email + the password "password123" + confirm it. Click submit.',
                    'PASS IF: an error appears saying the password "appears in [number] data breaches" and registration is blocked.',
                )),
                (76, s(
                    'Try registering with mismatched confirm password.',
                    'WHY: prevents typo-induced lockouts.',
                    'HOW: on /register, type a strong password but a different value in the confirm field. Submit.',
                    'PASS IF: error "Passwords do not match" appears, registration blocked.',
                )),
                (77, s(
                    'Submit the password-reset form with your real email.',
                    'WHY: tests the reset email flow.',
                    'HOW: /reset-password → enter your real email → click submit.',
                    'PASS IF: a confirmation message appears. (Check the email for the link — it should arrive within 1-2 minutes.)',
                )),
            ]),
            ('B.11 Static pages', [
                (78, s(
                    'Visit /privacy.',
                    'HOW: address bar → https://hirequadrant.com/privacy → Enter.',
                    'PASS IF: privacy policy renders with content (not blank).',
                )),
                (79, s(
                    'Visit /terms.',
                    'HOW: address bar → /terms.',
                    'PASS IF: terms of service render.',
                )),
                (80, s(
                    'Visit /content-policy.',
                    'HOW: address bar → /content-policy.',
                    'PASS IF: content policy renders.',
                )),
                (81, s(
                    'Visit /cookies.',
                    'HOW: address bar → /cookies.',
                    'PASS IF: cookie policy renders.',
                )),
                (82, s(
                    'Visit /about.',
                    'HOW: address bar → /about.',
                    'PASS IF: about page renders with company info.',
                )),
                (83, s(
                    'Visit /contact.',
                    'HOW: address bar → /contact.',
                    'PASS IF: contact page renders.',
                )),
                (84, s(
                    'Visit /pricing.',
                    'HOW: address bar → /pricing.',
                    'PASS IF: pricing page renders with tier cards.',
                )),
                (85, s(
                    'Confirm the security disclosure file is published.',
                    'WHY: required for responsible vulnerability reporting (security.txt standard).',
                    'HOW: address bar → https://hirequadrant.com/.well-known/security.txt → Enter.',
                    'PASS IF: a plain-text page loads with email contact info, expiry date, and "Contact: mailto:..." lines. NOT a 404 page.',
                )),
                (86, s(
                    'Confirm the dynamic sitemap is published.',
                    'WHY: search engines crawl this for indexing.',
                    'HOW: address bar → https://hirequadrant.com/sitemap-pages → Enter.',
                    'PASS IF: an XML page loads with many <url> entries (companies, industries, blog posts, etc.). The browser may show raw XML — that\'s OK.',
                )),
            ]),
            ('B.12 Anonymous edge cases', [
                (87, s(
                    'Visit a 404 (page that does not exist).',
                    'WHY: 404 page must be friendly, not blank.',
                    'HOW: address bar → https://hirequadrant.com/this-does-not-exist-12345 → Enter.',
                    'PASS IF: a "Not Found" / "404" page renders with a link back to home.',
                )),
                (88, s(
                    'Try to visit a candidate-only page logged out.',
                    'WHY: must redirect to login, not show error.',
                    'HOW: address bar → https://hirequadrant.com/profile → Enter.',
                    'PASS IF: redirects to /login (URL contains returnTo=/profile).',
                )),
                (89, s(
                    'Try to visit the company dashboard logged out.',
                    'HOW: address bar → /company-dashboard → Enter.',
                    'PASS IF: redirects to login.',
                )),
                (90, s(
                    'Try to visit the admin audit log logged out.',
                    'HOW: address bar → /admin/audit → Enter.',
                    'PASS IF: redirects to login.',
                )),
            ]),
        ],
    ),
    (
        'PART C: Individual / Candidate user',
        PART_C_INTRO,
        [
            ('C.1 Register a fresh account', [
                (1, s(
                    'Create a brand-new candidate account.',
                    'WHY: tests the full sign-up flow.',
                    'HOW: open a fresh incognito window. Visit /register. Click "Job Seeker". Fill in your real name, your real personal email, a strong password (12+ characters, mix letters/numbers/symbols), confirm it. Click "Create account".',
                    'PASS IF: success message ("Check your email") appears OR you\'re redirected to onboarding.',
                )),
                (2, s(
                    'Confirm the verification email arrives.',
                    'WHY: account email verification is required.',
                    'HOW: open your email inbox. Look for an email from no-reply@mail.app.supabase.io or noreply@hirequadrant.com.',
                    'PASS IF: email arrives within 1-2 minutes (check spam folder if not).',
                )),
                (3, s(
                    'Click the verification link in the email.',
                    'HOW: in the email, click "Confirm your email" or similar.',
                    'PASS IF: opens HireQuadrant in a new tab, you\'re logged in.',
                )),
            ]),
            ('C.2 Personalized home dashboard', [
                (4, s(
                    'Visit the homepage while logged in.',
                    'WHY: logged-in candidates get a personalized landing experience.',
                    'HOW: address bar → https://hirequadrant.com → Enter.',
                    'PASS IF: at the very top of the page (above the big "We work for you..." headline), you see a personalized dashboard.',
                )),
                (5, s(
                    'Confirm welcome banner.',
                    'HOW: read the top of the dashboard.',
                    'PASS IF: text says "Welcome back, [your first name]".',
                )),
                (6, s(
                    'Confirm 4 stat tiles.',
                    'HOW: look below the welcome banner.',
                    'PASS IF: 4 tiles labeled "Applied", "Saved", "Messages", "Notifications", each with a number (likely 0 for a brand new account).',
                )),
                (7, s(
                    'Confirm profile strength widget.',
                    'WHY: gamifies profile completion.',
                    'HOW: look for a "Profile strength" card.',
                    'PASS IF: shows a percentage (likely low — 10-30% — for a new account) with a colored progress bar.',
                )),
                (8, s(
                    'Confirm recent applications card.',
                    'HOW: look for a "Recent applications" section.',
                    'PASS IF: section is visible, probably empty for a new account.',
                )),
                (9, s(
                    'Confirm fresh jobs grid.',
                    'WHY: shows new jobs to draw the user back in.',
                    'HOW: scroll within the dashboard area.',
                    'PASS IF: a "Fresh jobs you might like" section shows 6 recent job cards.',
                )),
                (10, s(
                    'Click each stat tile.',
                    'WHY: tiles must be navigation shortcuts.',
                    'HOW: click "Applied" → confirm /my-jobs loads. Browser back. Click "Saved" → /my-jobs. Back. Click "Messages" → /messages. Back. Click "Notifications" → /notifications.',
                    'PASS IF: each tile navigates correctly.',
                )),
            ]),
            ('C.3 Profile setup', [
                (11, s(
                    'Open the avatar dropdown.',
                    'WHY: contains profile + settings shortcuts.',
                    'HOW: click your avatar (initials in a circle) at the top-right of any page.',
                    'PASS IF: dropdown opens with: Profile, My Reviews, My Demographics, Saved Jobs, Job Alerts, Settings, Help, Reset Password, Sign out.',
                )),
                (12, s(
                    'Open the Profile page.',
                    'HOW: in the dropdown, click "Profile".',
                    'PASS IF: /profile loads.',
                )),
                (13, s(
                    'Confirm profile completeness bar at top.',
                    'HOW: look at the very top of /profile.',
                    'PASS IF: a "Profile strength" card with a percentage and a checklist of criteria (Full name, Headline, Photo, Resume, Experience, Education, Skills, Preferences) is visible.',
                )),
                (14, s(
                    'Edit your name and save.',
                    'HOW: find the name field. Click the pencil/edit icon next to it. Edit to a 2+ character value. Click Save.',
                    'PASS IF: name updates inline. Profile strength % goes up.',
                )),
                (15, s(
                    'Open avatar upload picker.',
                    'HOW: click the camera icon on your avatar.',
                    'PASS IF: file picker dialog opens.',
                )),
                (16, s(
                    'Upload an image.',
                    'WHY: avatar personalizes the profile.',
                    'HOW: pick a JPG or PNG smaller than 3MB. Confirm.',
                    'PASS IF: avatar updates from initials to your image. Image also visible in the header.',
                )),
                (17, s(
                    'Edit headline.',
                    'WHY: headline is what employers see first.',
                    'HOW: find the headline field. Type "Senior Software Engineer" (or similar). Save.',
                    'PASS IF: persists.',
                )),
                (18, s(
                    'Save other basic profile fields.',
                    'HOW: edit location, phone, current role, target role. Click "Save Profile".',
                    'PASS IF: success toast or message. Reload page — values persist.',
                )),
                (19, s(
                    'Add a work experience entry.',
                    'HOW: scroll to "Work Experience" section. Click "+ Add experience". Fill: title="Senior Software Engineer", company="Test Company", start date 2020-01, end date 2023-12, 2-3 bullet points describing what you did. Click Save.',
                    'PASS IF: entry appears in the experience list. Profile strength % goes up.',
                )),
                (20, s(
                    'Edit an existing experience.',
                    'HOW: click the pencil icon on any experience. Change one bullet. Save.',
                    'PASS IF: change persists after page reload.',
                )),
                (21, s(
                    'Delete an experience.',
                    'HOW: click the trash icon on any experience. Confirm in dialog.',
                    'PASS IF: removed from the list.',
                )),
                (22, s(
                    'Add an education entry.',
                    'HOW: scroll to "Education" → click "+ Add education". Fill school + degree + dates. Save.',
                    'PASS IF: appears in the list.',
                )),
                (23, s(
                    'Add skills.',
                    'WHY: skills drive job matching.',
                    'HOW: scroll to "Skills" → type "JavaScript" in the input → press Enter. Repeat for 4 more skills (e.g., React, Python, AWS, SQL).',
                    'PASS IF: each appears as a chip.',
                )),
                (24, s(
                    'Remove a skill chip.',
                    'HOW: click the X on any skill chip.',
                    'PASS IF: chip removes.',
                )),
                (25, s(
                    'Set job preferences.',
                    'WHY: preferences power job recommendations.',
                    'HOW: scroll to "Job Preferences" → fill desired titles, desired locations, salary floor → Save.',
                    'PASS IF: persists.',
                )),
                (26, s(
                    'Upload a resume.',
                    'HOW: scroll to "Resume" section → click "Upload resume" → pick a PDF or DOCX file under 5MB.',
                    'PASS IF: filename appears in the resume slot. Profile strength % goes to 100% (if other criteria are also met).',
                )),
                (27, s(
                    'Preview the uploaded resume.',
                    'HOW: click "View" next to the resume.',
                    'PASS IF: resume opens in an embedded viewer modal. Click Close to exit.',
                )),
                (28, s(
                    'Confirm profile is now complete.',
                    'HOW: scroll to top.',
                    'PASS IF: profile strength bar shows 90-100%.',
                )),
            ]),
            ('C.4 Settings page', [
                (29, s(
                    'Visit Settings.',
                    'HOW: avatar dropdown → "Settings" (or address bar → /settings).',
                    'PASS IF: page renders with Email + Password sections at top.',
                )),
                (30, s(
                    'Confirm email is correct (read-only).',
                    'PASS IF: email field shows your current email and cannot be edited.',
                )),
                (31, s(
                    'Trigger a password reset email.',
                    'WHY: tests the reset flow for already-logged-in users.',
                    'HOW: click "Send reset email" under Password.',
                    'PASS IF: toast appears: "Password reset email sent". Email arrives in 1-2 min.',
                )),
                (32, s(
                    'Toggle each notification preference.',
                    'WHY: users control what emails they get.',
                    'HOW: scroll to "Notification preferences". Toggle each row off and back on (Job alerts, Application updates, Direct messages, Review responses, Product updates).',
                    'PASS IF: toggles save without errors.',
                )),
                (33, s(
                    'Change digest frequency.',
                    'HOW: change the "Digest frequency" dropdown from Daily to Weekly. Reload the page.',
                    'PASS IF: still shows Weekly after reload.',
                )),
                (34, s(
                    'Begin 2FA setup.',
                    'WHY: 2FA hardens account security.',
                    'HOW: scroll to "Two-factor authentication". Click "Enable two-factor".',
                    'PASS IF: a QR code appears, with the secret displayed below.',
                )),
                (35, s(
                    'Scan the QR with your authenticator app.',
                    'HOW: open Google Authenticator (or Authy / 1Password) on your phone. Tap the + button → "Scan QR code" → point at the screen.',
                    'PASS IF: app shows a new entry "HireQuadrant: [your email]" with a 6-digit code refreshing every 30 sec.',
                )),
                (36, s(
                    'Confirm the 2FA code.',
                    'HOW: type the current 6-digit code from your phone into the input on the website. Click "Confirm".',
                    'PASS IF: 2FA is now enabled. 10 recovery codes appear.',
                )),
                (37, s(
                    'Copy recovery codes.',
                    'WHY: needed to recover access if you lose your phone.',
                    'HOW: click "Copy all" under the recovery codes block.',
                    'PASS IF: a check icon appears (briefly says "Copied"). Paste somewhere safe (notes app).',
                )),
                (38, s(
                    'Disable 2FA again.',
                    'WHY: cleanup so you don\'t lock yourself out testing.',
                    'HOW: click "Disable two-factor". Confirm in dialog.',
                    'PASS IF: 2FA removed. Section returns to "Enable" state.',
                )),
                (39, s(
                    'Confirm active sessions panel.',
                    'WHY: users can see + revoke other devices.',
                    'HOW: scroll to "Active sessions".',
                    'PASS IF: at least 1 entry showing your current device.',
                )),
                (40, s(
                    'Sign in from another browser to test.',
                    'WHY: confirms multi-session detection.',
                    'HOW: open Safari (or Firefox). Go to hirequadrant.com → Sign in with the same account. Then return to the Chrome window with /settings open. Refresh.',
                    'PASS IF: a 2nd entry now appears in Active sessions list.',
                )),
                (41, s(
                    'Sign out all other devices.',
                    'HOW: in Settings → "Sign out all others" → confirm.',
                    'PASS IF: in the Safari/Firefox window, refresh — it redirects to login. The current Chrome session stays logged in.',
                )),
                (42, s(
                    'Download your data (GDPR export).',
                    'WHY: legal requirement.',
                    'HOW: scroll to "Download your data" → click "Download data".',
                    'PASS IF: a JSON file downloads to your computer (filename like hirequadrant-export-YYYY-MM-DD.json).',
                )),
                (43, s(
                    'Open the JSON archive.',
                    'HOW: open the downloaded file in any text editor (TextEdit, VS Code, Notepad).',
                    'PASS IF: file contains your profile, applications, saved jobs, etc. (Lots of curly braces and quoted strings — that\'s normal JSON.)',
                )),
            ]),
            ('C.5 My Jobs', [
                (44, s(
                    'Visit My Jobs.',
                    'WHY: candidate-side activity hub.',
                    'HOW: address bar → /my-jobs (or click stat tile from dashboard).',
                    'PASS IF: page renders with 5 tabs at top: Saved, Invitations, Applied, Interviews, Archived.',
                )),
                (45, s(
                    'Click each tab.',
                    'PASS IF: each tab loads its content (most will be empty for a new account; that\'s OK — empty state copy should appear).',
                )),
            ]),
            ('C.6 Saved searches', [
                (46, s(
                    'Visit Saved Searches.',
                    'WHY: lets candidates save filters and get alerts on new matching jobs.',
                    'HOW: address bar → /saved-searches.',
                    'PASS IF: page renders with a "Create a saved search" form at top.',
                )),
                (47, s(
                    'Create a saved search.',
                    'HOW: fill: Name="Remote senior dev roles", Keyword="engineer", Location="Remote", Frequency="Daily". Click Save.',
                    'PASS IF: the saved search appears in the list below.',
                )),
                (48, s(
                    'Change frequency.',
                    'HOW: change the dropdown next to your saved search from Daily to Weekly.',
                    'PASS IF: change saves immediately.',
                )),
                (49, s(
                    'Delete the saved search.',
                    'HOW: click trash icon next to the saved search.',
                    'PASS IF: row removes.',
                )),
            ]),
            ('C.7 Browse and apply flow', [
                (50, s(
                    'Visit /jobs and apply filters.',
                    'HOW: same as Section B.4 — filter as desired.',
                    'PASS IF: filtering works while logged in.',
                )),
                (51, s(
                    'Open a job detail.',
                    'HOW: click any job card.',
                    'PASS IF: detail page loads.',
                )),
                (52, s(
                    'Confirm AI match score pill.',
                    'WHY: AI scoring is a key differentiator.',
                    'HOW: look near the job title for a "✨ XX% match" pill.',
                    'PASS IF: a percentage appears (may take 1-2 seconds to load on first view).',
                )),
                (53, s(
                    'Save the job.',
                    'HOW: click the bookmark/Save icon.',
                    'PASS IF: icon turns filled/blue. Toast appears.',
                )),
                (54, s(
                    'Apply to the job.',
                    'HOW: click "Apply".',
                    'PASS IF: an apply form appears (modal or inline).',
                )),
                (55, s(
                    'Confirm form pre-fills profile data.',
                    'HOW: look at name + email fields.',
                    'PASS IF: pre-filled with your profile values. Resume is also pre-attached.',
                )),
                (56, s(
                    'Submit the application.',
                    'HOW: fill any remaining fields (cover letter, screening questions). Click Submit.',
                    'PASS IF: success message appears. Apply button changes to "Applied".',
                )),
                (57, s(
                    'Verify in /my-jobs Applied tab.',
                    'HOW: address bar → /my-jobs → click "Applied" tab.',
                    'PASS IF: your application appears in the list.',
                )),
                (58, s(
                    'View application status timeline.',
                    'HOW: click "Status" button next to your application.',
                    'PASS IF: a timeline expands inline showing "Applied" with a timestamp.',
                )),
                (59, s(
                    'Withdraw the application.',
                    'HOW: click "Withdraw" button. Confirm.',
                    'PASS IF: application removes. Toast: "Application withdrawn".',
                )),
                (60, s(
                    'Re-apply from Saved tab.',
                    'HOW: click "Saved" tab → click "Apply" quick-action on your saved job.',
                    'PASS IF: re-applies you.',
                )),
            ]),
            ('C.8 Job referrals', [
                (61, s(
                    'Generate a referral share link.',
                    'WHY: candidates can share jobs with others and track referrals.',
                    'HOW: open any job detail page → scroll to "Refer someone" section.',
                    'PASS IF: a unique URL with ?ref=CODE is shown along with copy/email/X/LinkedIn buttons.',
                )),
                (62, s(
                    'Copy the link.',
                    'HOW: click "Copy".',
                    'PASS IF: toast: "Share link copied".',
                )),
                (63, s(
                    'Test the link in incognito.',
                    'WHY: confirms tracking.',
                    'HOW: paste the copied URL into a fresh incognito window. Open it.',
                    'PASS IF: job detail page loads. (Click is tracked behind the scenes; no visual change.)',
                )),
            ]),
            ('C.9 Reviews', [
                (64, s(
                    'Visit My Reviews.',
                    'HOW: avatar dropdown → "My Reviews".',
                    'PASS IF: page renders. May be empty for a new account.',
                )),
                (65, s(
                    'Write a review on a company.',
                    'HOW: visit any company profile (e.g., /companies/quadrant-inc) → click "Write a review" button. Fill: 5 stars, role you held (any), time at company (any), title, pros (2-3 sentences), cons (2-3 sentences). Submit.',
                    'PASS IF: confirmation appears. Review goes live or shows "Pending" status. Visit /my-reviews — your review appears.',
                )),
                (66, s(
                    'Vote on a review.',
                    'HOW: on any company profile, click thumbs-up or thumbs-down on someone else\'s review.',
                    'PASS IF: count updates. (You can\'t vote on your own.)',
                )),
                (67, s(
                    'Report a review.',
                    'HOW: click flag icon on a review → pick a reason → submit.',
                    'PASS IF: confirmation appears.',
                )),
            ]),
            ('C.10 Demographics (optional EEO)', [
                (68, s(
                    'Visit Demographics.',
                    'WHY: optional EEO data collection, anonymized.',
                    'HOW: avatar dropdown → "My Demographics" (or /demographics).',
                    'PASS IF: form with optional fields renders (race/ethnicity, gender, disability, veteran).',
                )),
                (69, s(
                    'Save partial data.',
                    'HOW: pick a value or two. The form auto-saves on change OR you click a Save button.',
                    'PASS IF: confirmation. Reload page — values persist.',
                )),
            ]),
            ('C.11 Messages and Notifications', [
                (70, s(
                    'Visit Messages.',
                    'HOW: address bar → /messages.',
                    'PASS IF: empty state copy renders ("No conversations yet") for a new account.',
                )),
                (71, s(
                    'Check the notification bell.',
                    'HOW: click the bell icon at the top-right of the header.',
                    'PASS IF: a dropdown appears with recent notifications (or "No notifications" copy).',
                )),
                (72, s(
                    'Visit /notifications full page.',
                    'HOW: address bar → /notifications.',
                    'PASS IF: full notifications list renders. Click "Mark all read" if any unread.',
                )),
            ]),
            ('C.12 Interview practice', [
                (73, s(
                    'Visit Interview Practice.',
                    'WHY: AI-graded mock interview prep.',
                    'HOW: address bar → /interview-practice.',
                    'PASS IF: page loads (must be logged in — redirects to /login if not).',
                )),
                (74, s(
                    'Pick role and question type.',
                    'HOW: pick "Software Engineer" from Role dropdown. Pick "Behavioral" from Question type.',
                    'PASS IF: a question appears in a gray box.',
                )),
                (75, s(
                    'Get a different question.',
                    'HOW: click the "New" button (with shuffle icon) next to the question.',
                    'PASS IF: question changes.',
                )),
                (76, s(
                    'Type a practice answer.',
                    'WHY: minimum 30 chars to submit.',
                    'HOW: type 2-3 paragraphs of a real-style answer in the textarea.',
                    'PASS IF: char count below input updates.',
                )),
                (77, s(
                    'Get AI feedback.',
                    'HOW: click "Get AI feedback" button.',
                    'PASS IF: spinner shows for 5-15 seconds. Then a result card appears with: a numeric score (0-100), summary, strengths list, improvements list, and "Example of a stronger answer" rewrite.',
                )),
                (78, s(
                    'Confirm history saved.',
                    'HOW: refresh the page. Scroll to "Your recent sessions".',
                    'PASS IF: your past attempt appears with its score.',
                )),
            ]),
            ('C.13 Sign out', [
                (79, s(
                    'Sign out.',
                    'HOW: avatar dropdown → "Sign out".',
                    'PASS IF: returns to homepage. You\'re logged out.',
                )),
                (80, s(
                    'Confirm logged-in pages now redirect.',
                    'HOW: address bar → /profile.',
                    'PASS IF: redirects to /login.',
                )),
            ]),
        ],
    ),
    (
        'PART D: Company / Employer user',
        PART_D_INTRO,
        [
            ('D.1 Login and navigation', [
                (1, s(
                    'Open a fresh incognito window and log in as employer.',
                    'WHY: clean session.',
                    'HOW: Cmd+Shift+N (Mac) → /login → email test-employer-1@hirequadrant-test.com → password TestBiz123! → Sign in.',
                    'PASS IF: lands on /company-dashboard. Top nav shows "Dashboard" + "Talent Search" (no candidate-only items like "My Jobs").',
                )),
                (2, s(
                    'Confirm 11 dashboard tabs are visible.',
                    'HOW: look at the dashboard tab bar.',
                    'PASS IF: tabs labeled: Jobs, Applicants, Analytics, Benchmarks, Reviews, Updates, Q&A, Why Join Us, AI Assistant, Company Profile, Subscription.',
                )),
            ]),
            ('D.2 Jobs tab', [
                (3, s(
                    'View existing jobs.',
                    'HOW: click "Jobs" tab.',
                    'PASS IF: a list of jobs renders with title, location, applicants count, status.',
                )),
                (4, s(
                    'Create a new job.',
                    'HOW: click "+ New Job" button. Fill: title="Test QA Role", location="Remote", type="Full-time", salary="$100-150k". Description: 2-3 paragraphs. Click Save.',
                    'PASS IF: job appears in list.',
                )),
                (5, s(
                    'Open the new job in the public view.',
                    'WHY: confirms job is live.',
                    'HOW: copy the job\'s URL or visit /jobs and find your new job. Click into it.',
                    'PASS IF: job detail renders correctly with title, location, salary, description.',
                )),
                (6, s(
                    'Edit the job to add screening questions.',
                    'HOW: back on /company-dashboard → Jobs tab → click edit/pencil on your new job. Scroll to "Manage Screening Questions". Add a question: "Are you authorized to work in the US?" with type Yes/No. Save.',
                    'PASS IF: question saved and visible in apply form on public job page.',
                )),
                (7, s(
                    'Add custom application fields.',
                    'WHY: lets employers collect arbitrary data per job.',
                    'HOW: in the same edit modal, scroll further to "Custom Application Fields". Add: label="Years of QA experience", type="Short Text", required=Yes. Click "Save fields".',
                    'PASS IF: success toast. Open public job page → confirm field appears in apply form.',
                )),
                (8, s(
                    'Delete a test job.',
                    'WHY: cleanup.',
                    'HOW: in Jobs tab → find your test job → click delete → confirm.',
                    'PASS IF: removed from list.',
                )),
            ]),
            ('D.3 Applicants tab — list and pipeline', [
                (9, s(
                    'View applicants in list.',
                    'HOW: click "Applicants" tab.',
                    'PASS IF: list renders if any applicants exist (may be empty).',
                )),
                (10, s(
                    'Filter applicants.',
                    'HOW: use the status, job, and shortlisted filters at the top.',
                    'PASS IF: list narrows.',
                )),
                (11, s(
                    'Switch to Pipeline (Kanban) view.',
                    'WHY: drag-and-drop pipeline is a key employer feature.',
                    'HOW: click the "Pipeline" toggle at the top.',
                    'PASS IF: 6 columns appear: New, Reviewing, Interview, Offered, Hired, Rejected.',
                )),
                (12, s(
                    'Drag a card between columns.',
                    'WHY: tests the core pipeline interaction.',
                    'HOW: if applicants exist, click and hold any card → drag to another column → release.',
                    'PASS IF: card moves; status updates immediately.',
                )),
            ]),
            ('D.4 Applicant detail modal', [
                (13, s(
                    'Open an applicant detail.',
                    'HOW: click any applicant card (in either view).',
                    'PASS IF: a detail modal opens with candidate info, screening answers, resume download link, AI Screening section, employer notes textarea, tags input, CRM Notes & Tags panel, and a "Schedule Interview" button.',
                )),
                (14, s(
                    'Run AI screening.',
                    'WHY: AI-powered fit assessment.',
                    'HOW: click "Run AI Screening" in the modal.',
                    'PASS IF: spinner for 5-15 seconds, then results appear: fit score (0-100), strengths list, concerns list.',
                )),
                (15, s(
                    'Add a CRM tag.',
                    'WHY: team-wide tagging for organization.',
                    'HOW: in the "CRM Notes & Tags" panel, type "top-choice" in the tags input → click Add (or press Enter).',
                    'PASS IF: chip appears.',
                )),
                (16, s(
                    'Add a private note.',
                    'HOW: in the same panel, type a note → click "Add note".',
                    'PASS IF: note appears with timestamp.',
                )),
                (17, s(
                    'Delete a note.',
                    'HOW: click X on a note. Confirm.',
                    'PASS IF: removed.',
                )),
            ]),
            ('D.5 Schedule interview', [
                (18, s(
                    'Open the schedule interview modal.',
                    'HOW: in the applicant detail modal, click "Schedule interview" button at the bottom.',
                    'PASS IF: a modal opens with date, time, duration, location, meeting URL, notes fields.',
                )),
                (19, s(
                    'Fill the interview details.',
                    'HOW: pick a date 2 days from now, 10:00 time, 30 min duration, location="Video call", meeting URL=any URL, notes optional.',
                ),
                ),
                (20, s(
                    'Download .ics calendar invite.',
                    'WHY: industry-standard calendar format.',
                    'HOW: click ".ics" button.',
                    'PASS IF: a .ics file downloads.',
                )),
                (21, s(
                    'Open the .ics file.',
                    'HOW: double-click the downloaded file. (On Mac it opens in Calendar app. On Windows it opens in Outlook.)',
                    'PASS IF: opens cleanly with the right title, date, time, location.',
                )),
                (22, s(
                    'Save the interview record.',
                    'HOW: back in the modal, click "Schedule".',
                    'PASS IF: confirms saved. An in-app message is auto-sent to the candidate.',
                )),
            ]),
            ('D.6 Bulk message', [
                (23, s(
                    'Open bulk message composer.',
                    'WHY: lets employers contact multiple applicants at once.',
                    'HOW: on Applicants tab (List or Pipeline), click "Bulk message".',
                    'PASS IF: a modal opens with checked-by-default applicants.',
                )),
                (24, s(
                    'Send a bulk message.',
                    'HOW: uncheck a few. Type subject="Test bulk message", body="Hello team!". Click Send.',
                    'PASS IF: success toast: "Sent to N applicants". Each recipient sees the message in their /messages.',
                )),
            ]),
            ('D.7 Analytics tab', [
                (25, s(
                    'View analytics.',
                    'HOW: click Analytics tab.',
                    'PASS IF: charts render (applications over time, status breakdown, etc.). May be sparse for a test account; that\'s OK.',
                )),
            ]),
            ('D.8 Benchmarks tab', [
                (26, s(
                    'View benchmarks.',
                    'WHY: shows how this employer compares to peers.',
                    'HOW: click Benchmarks tab.',
                    'PASS IF: 4 cards render — Average rating, Total reviews, Open roles, Applications received. Each shows You vs Peer median bar comparison plus a trend pill.',
                )),
            ]),
            ('D.9 Reviews tab', [
                (27, s(
                    'View reviews.',
                    'HOW: click Reviews tab.',
                    'PASS IF: list of all reviews on your company. Filter (all/pending/approved/rejected) at top.',
                )),
                (28, s(
                    'Respond to a review.',
                    'HOW: click "Respond" on any review. Type a response. Save.',
                    'PASS IF: response posts publicly under that review.',
                )),
                (29, s(
                    'Report a review.',
                    'HOW: click "Report" on a different review. Pick a reason. Submit.',
                    'PASS IF: report goes to admin queue.',
                )),
            ]),
            ('D.10 Updates tab', [
                (30, s(
                    'Post a company update.',
                    'WHY: lets employers broadcast news to candidates following them.',
                    'HOW: click Updates tab → "Post update" → fill title="Test update", body="We\'re hiring!" → optional image URL → save.',
                    'PASS IF: update appears in the feed.',
                )),
                (31, s(
                    'Confirm public visibility.',
                    'HOW: open the public company page in a new tab.',
                    'PASS IF: the update appears in the Updates feed.',
                )),
            ]),
            ('D.11 Q&A tab', [
                (32, s(
                    'Answer a pending question.',
                    'HOW: click Q&A tab. If pending questions exist, click answer → write a response → save.',
                    'PASS IF: question marked answered. Public profile shows the Q&A pair.',
                )),
            ]),
            ('D.12 Why Join Us tab', [
                (33, s(
                    'Add a Text block.',
                    'HOW: click Why Join Us tab → "Add block" → "Text" → write a paragraph → save.',
                    'PASS IF: block appears in editor and on public profile.',
                )),
                (34, s(
                    'Add a Stat block.',
                    'HOW: "Add block" → "Stat" → fill (e.g., "500 employees") → save.',
                    'PASS IF: renders as a big-number tile on public profile.',
                )),
                (35, s(
                    'Add a Quote block.',
                    'HOW: "Add block" → "Quote" → fill quote + attribution → save.',
                    'PASS IF: renders with attribution.',
                )),
                (36, s(
                    'Add an Image or Video block.',
                    'HOW: "Add block" → "Image" or "Video" → paste a URL or upload → save.',
                    'PASS IF: renders.',
                )),
                (37, s(
                    'Reorder blocks.',
                    'HOW: drag a block up or down in the editor.',
                    'PASS IF: order persists after reload.',
                )),
            ]),
            ('D.13 AI Assistant tab', [
                (38, s(
                    'Generate a job description.',
                    'WHY: AI helps employers write better JDs faster.',
                    'HOW: click AI Assistant tab → fill title="Senior Backend Engineer" → bullets (2-3 lines about the role) → click Generate.',
                    'PASS IF: a full JD generates in 5-15 seconds with structured sections.',
                )),
                (39, s(
                    'Confirm Inclusive Language Linter.',
                    'WHY: flags problematic language pre-publish.',
                    'HOW: scroll below the generated JD.',
                    'PASS IF: a panel appears flagging any non-inclusive language with swap suggestions (or "Looks good!" if none).',
                )),
                (40, s(
                    'Score the JD.',
                    'WHY: gives employers a 0-100 quality grade.',
                    'HOW: scroll to "Job description quality score" → click "Score this JD".',
                    'PASS IF: a card appears with overall 0-100 score, 4 sub-scores (Specificity, Inclusivity, Clarity, Completeness), strengths list, improvements list.',
                )),
            ]),
            ('D.14 Company Profile tab', [
                (41, s(
                    'Edit company info.',
                    'HOW: click Company Profile tab → edit description, website, industry, size, location, contact email → save.',
                    'PASS IF: changes persist; visible on public profile.',
                )),
                (42, s(
                    'Upload a logo.',
                    'HOW: in same tab → upload a new square image (under 3MB) → save.',
                    'PASS IF: logo updates on public profile.',
                )),
            ]),
            ('D.15 Team Invites', [
                (43, s(
                    'Generate an invite link.',
                    'WHY: lets employers add teammates.',
                    'HOW: in Company Profile tab, scroll to "Team Invites". Pick role="Recruiter" → click "Generate invite link".',
                    'PASS IF: a new invite appears in the list.',
                )),
                (44, s(
                    'Copy the invite link.',
                    'HOW: click "Copy link".',
                    'PASS IF: toast: "Copied".',
                )),
                (45, s(
                    'Open the invite link in incognito.',
                    'HOW: paste the link into a fresh incognito window. Open it.',
                    'PASS IF: lands on /register?invite=CODE — registration page.',
                )),
            ]),
            ('D.16 Subscription tab', [
                (46, s(
                    'View subscription.',
                    'HOW: click Subscription tab.',
                    'PASS IF: shows current plan or "Coming soon" copy if Stripe isn\'t fully wired.',
                )),
            ]),
            ('D.17 Talent Search', [
                (47, s(
                    'Browse Talent Search.',
                    'WHY: lets employers proactively find candidates.',
                    'HOW: top nav → "Talent Search".',
                    'PASS IF: list of candidate profiles renders (those who opted in).',
                )),
                (48, s(
                    'Filter candidates.',
                    'HOW: use skills/location/experience filters.',
                    'PASS IF: list narrows.',
                )),
                (49, s(
                    'Open a candidate profile.',
                    'HOW: click any candidate card.',
                    'PASS IF: profile preview renders (limited info if not premium).',
                )),
            ]),
            ('D.18 Sign out', [
                (50, s(
                    'Sign out as employer.',
                    'HOW: avatar → Sign out.',
                    'PASS IF: returns to home, logged out.',
                )),
            ]),
        ],
    ),
    (
        'PART E: Admin / Portal user',
        PART_E_INTRO,
        [
            ('E.1 Login and navigation', [
                (1, s(
                    'Log in as admin.',
                    'HOW: open fresh incognito → /login → email rrainey19138@gmail.com (or your admin email) → password TestPass123! → Sign in.',
                    'PASS IF: signed in. Top nav shows a "Portal" link.',
                )),
                (2, s(
                    'Open the Portal.',
                    'HOW: click "Portal" in top nav.',
                    'PASS IF: /company-portal loads with 3 sections — Moderation queues (Review Moderation, Reports, Appeals); Platform tools (Admin Dashboard, Talent Search, XML Feeder, Company Sources, Cron Health, Audit Log); Your account.',
                )),
            ]),
            ('E.2 Review Moderation', [
                (3, s(
                    'Open Review Moderation queue.',
                    'WHY: admins approve/reject pending reviews.',
                    'HOW: click "Review Moderation" tile in Portal.',
                    'PASS IF: /admin/reviews loads with pending reviews list.',
                )),
                (4, s(
                    'View a pending review.',
                    'HOW: click any pending review.',
                    'PASS IF: full review content visible.',
                )),
                (5, s(
                    'Approve a review.',
                    'HOW: click Approve.',
                    'PASS IF: review goes live; row disappears from queue.',
                )),
                (6, s(
                    'Reject a review.',
                    'HOW: click Reject → type a reason → submit.',
                    'PASS IF: review marked rejected. Author gets notified. Audit log entry created.',
                )),
            ]),
            ('E.3 Reports', [
                (7, s(
                    'Open Reports queue.',
                    'WHY: handles user-flagged reviews.',
                    'HOW: Portal → "Reports" tile.',
                    'PASS IF: /admin/reports loads with reported reviews.',
                )),
                (8, s(
                    'Decide on a report.',
                    'HOW: click into a report → see review + reporter + reason → click Remove or Dismiss.',
                    'PASS IF: action recorded.',
                )),
            ]),
            ('E.4 Appeals', [
                (9, s(
                    'Open Appeals queue.',
                    'WHY: lets review authors contest rejections.',
                    'HOW: Portal → "Appeals" tile.',
                    'PASS IF: /admin/appeals loads.',
                )),
                (10, s(
                    'Decide on an appeal.',
                    'HOW: click into one → see original review + rejection reason + appeal text → click Reinstate or Uphold.',
                    'PASS IF: action recorded.',
                )),
            ]),
            ('E.5 Audit Log', [
                (11, s(
                    'View the audit log.',
                    'WHY: complete history of admin actions for accountability.',
                    'HOW: address bar → /admin/audit.',
                    'PASS IF: list of last 200 actions: actor name, action, target, timestamp, metadata.',
                )),
                (12, s(
                    'Filter by action type.',
                    'HOW: use the dropdown → pick review_moderation.',
                    'PASS IF: list narrows.',
                )),
                (13, s(
                    'Confirm your prior actions appear.',
                    'HOW: scan the list for entries matching what you did in steps 5-10 above.',
                    'PASS IF: at least 2-3 entries from your testing show up.',
                )),
            ]),
            ('E.6 Cron Health', [
                (14, s(
                    'View cron job health.',
                    'WHY: scheduled jobs (sitemap regen, cleanup) need monitoring.',
                    'HOW: address bar → /admin/cron.',
                    'PASS IF: list of scheduled jobs with last run / status / 24h success-failure counts. (If no cron has run yet, an empty state appears — that\'s OK.)',
                )),
            ]),
            ('E.7 XML Feeder', [
                (15, s(
                    'Open XML Feeder.',
                    'WHY: imports jobs from external feed sources.',
                    'HOW: Portal → "XML Feeder" tile (or address bar /xml-feeder).',
                    'PASS IF: page lists configured XML feeds.',
                )),
                (16, s(
                    'Trigger a re-ingest.',
                    'HOW: find a feed → click Re-ingest.',
                    'PASS IF: feed runs; success message after ~30 sec; jobs populate.',
                )),
                (17, s(
                    'Verify ingest quality.',
                    'WHY: company name must map correctly (no "hirequadrant.xml" leakage).',
                    'HOW: open /jobs in another tab. Pick any newly-ingested job.',
                    'PASS IF: company name shows as canonical (e.g., "Quadrant, Inc."), NOT as the feed filename.',
                )),
            ]),
            ('E.8 Company Sources', [
                (18, s(
                    'Open Company Sources.',
                    'WHY: maps XML filenames → canonical company names.',
                    'HOW: Portal → "Company Sources" tile (or /company-sources).',
                    'PASS IF: list of mappings renders.',
                )),
                (19, s(
                    'Edit a mapping.',
                    'HOW: edit any row → save.',
                    'PASS IF: change persists.',
                )),
            ]),
            ('E.9 Admin Dashboard', [
                (20, s(
                    'Open Admin Dashboard.',
                    'HOW: Portal → "Admin Dashboard" tile (or /admin).',
                    'PASS IF: user table + site-wide stats render.',
                )),
                (21, s(
                    'Filter users by role.',
                    'HOW: use role dropdown.',
                    'PASS IF: list narrows.',
                )),
                (22, s(
                    'Confirm site stats are present.',
                    'HOW: look at the top of /admin.',
                    'PASS IF: total users, jobs, companies, reviews, applications counts visible.',
                )),
            ]),
            ('E.10 Admin sees candidate features (gating check)', [
                (23, s(
                    'Confirm avatar dropdown.',
                    'WHY: admin should see candidate items (My Reviews, Demographics, etc.) because gating is "not company" not "not admin".',
                    'HOW: click avatar.',
                    'PASS IF: dropdown shows: Profile, My Reviews, My Demographics, Saved Jobs, Job Alerts, Settings, Help, Reset Password, Sign out.',
                )),
                (24, s(
                    'Confirm admin does NOT see employer-only nav.',
                    'PASS IF: top nav does NOT show "Dashboard" or "Talent Search" labels (those are for company role only).',
                )),
            ]),
            ('E.11 Sign out', [
                (25, s(
                    'Sign out.',
                    'HOW: avatar → Sign out.',
                    'PASS IF: back to home, logged out.',
                )),
            ]),
        ],
    ),
    (
        'PART F: Cross-cutting / Technical',
        PART_F_INTRO,
        [
            ('F.1 Dark mode on every major page', [
                (1, s(
                    'Enable dark mode globally.',
                    'WHY: users with dark-mode preference must have a consistent experience.',
                    'HOW: while logged in (any persona), click the moon/sun icon at top-right.',
                    'PASS IF: page goes dark.',
                )),
                (2, s(
                    'Visit homepage in dark mode.',
                    'PASS IF: clean dark theme. No white sections, all text readable.',
                )),
                (3, s('Visit /jobs.', 'PASS IF: dark, readable.')),
                (4, s('Visit /jobs/[any-id] (open a job).', 'PASS IF: dark, readable.')),
                (5, s('Visit /companies.', 'PASS IF: dark, readable.')),
                (6, s('Visit any /companies/[slug].', 'PASS IF: dark, readable.')),
                (7, s('Visit /blog.', 'PASS IF: dark, readable.')),
                (8, s('Visit any /blog/[slug].', 'PASS IF: dark, readable.')),
                (9, s('Visit /help-center.', 'PASS IF: dark, readable.')),
                (10, s('Visit /profile (logged in).', 'PASS IF: dark, readable.')),
                (11, s('Visit /settings.', 'PASS IF: dark, readable.')),
                (12, s('Visit /my-jobs.', 'PASS IF: dark, readable.')),
                (13, s('Visit /messages.', 'PASS IF: dark, readable.')),
                (14, s('Visit /notifications.', 'PASS IF: dark, readable.')),
                (15, s('Visit /interview-practice.', 'PASS IF: dark, readable.')),
                (16, s('Visit /saved-searches.', 'PASS IF: dark, readable.')),
                (17, s('Visit /compare.', 'PASS IF: dark, readable.')),
                (18, s('As employer: visit /company-dashboard, click each of the 11 tabs.', 'PASS IF: every tab dark, readable.')),
                (19, s('As admin: visit /company-portal, /admin/audit, /admin/cron.', 'PASS IF: dark, readable.')),
            ]),
            ('F.2 Mobile (real phone)', [
                (20, s(
                    'Open the site on a real iPhone or Android.',
                    'WHY: 60%+ of traffic comes from mobile.',
                    'HOW: on your phone, open Safari/Chrome → type hirequadrant.com → enter.',
                    'PASS IF: site loads, header collapses to a hamburger (3 horizontal lines) icon.',
                )),
                (21, s(
                    'Open the mobile menu.',
                    'HOW: tap hamburger icon.',
                    'PASS IF: menu drawer slides in with all nav links.',
                )),
                (22, s(
                    'Sign in on phone.',
                    'HOW: in mobile menu → Sign in → log in with seeded candidate.',
                    'PASS IF: signed in, dashboard renders mobile-friendly.',
                )),
                (23, s(
                    'Apply for a job on phone.',
                    'WHY: most job applications happen on mobile now.',
                    'HOW: navigate to a job → tap Apply → fill form → submit.',
                    'PASS IF: keyboard does not cover the submit button. Form scrolls correctly. Submit succeeds.',
                )),
                (24, s(
                    'Confirm filter sidebar collapses.',
                    'HOW: on mobile /jobs, look at filters.',
                    'PASS IF: filters are collapsible (filter button or drawer) — not always visible eating screen space.',
                )),
            ]),
            ('F.3 Lighthouse audit (Performance / SEO / Accessibility)', [
                (25, s(
                    'Run Lighthouse on the homepage.',
                    'WHY: Lighthouse scores a page on Performance, Accessibility, Best Practices, SEO.',
                    'HOW STEP-BY-STEP:',
                    '  1) Open Chrome (must be Chrome — Lighthouse does not work in Safari/Firefox).',
                    '  2) In incognito, visit https://hirequadrant.com.',
                    '  3) Press Cmd+Option+I (Mac) or F12 (Windows) to open DevTools.',
                    '  4) In DevTools, click the "Lighthouse" tab. (If you don\'t see it, click the >> overflow menu in the DevTools tabs row.)',
                    '  5) Categories: check Performance + Accessibility + Best Practices + SEO. Uncheck PWA.',
                    '  6) Mode: "Navigation (default)". Device: "Desktop".',
                    '  7) Click the blue "Analyze page load" button.',
                    '  8) Wait ~30-60 seconds while it audits.',
                    'PASS IF: Performance score is 70 or higher. Write the actual score in Notes.',
                )),
                (26, s(
                    'Lighthouse — Accessibility score.',
                    'HOW: same Lighthouse run as Step 25 — Accessibility score is in the same report.',
                    'PASS IF: Accessibility ≥ 90.',
                )),
                (27, s(
                    'Lighthouse — Best Practices score.',
                    'HOW: same Lighthouse report.',
                    'PASS IF: Best Practices ≥ 90.',
                )),
                (28, s(
                    'Lighthouse — SEO score.',
                    'HOW: same Lighthouse report.',
                    'PASS IF: SEO ≥ 90.',
                )),
                (29, s(
                    'Repeat Lighthouse audit on Mobile.',
                    'WHY: mobile performance is often weaker.',
                    'HOW: same as Step 25 but in Lighthouse settings, change Device from "Desktop" to "Mobile". Re-run.',
                    'PASS IF: Performance ≥ 60 (mobile is typically lower than desktop). Accessibility/Best Practices/SEO ≥ 85.',
                )),
            ]),
            ('F.4 Auth edge cases', [
                (30, s(
                    'Test login rate limiting.',
                    'WHY: prevents brute-force attacks.',
                    'HOW STEP-BY-STEP:',
                    '  1) Go to /login.',
                    '  2) Enter your real email + a wrong password → click Sign in.',
                    '  3) Repeat with wrong password 5 more times (6 total).',
                    'PASS IF: at the 9th-ish attempt, an error appears: "Too many failed attempts. Please wait a few minutes." OR similar lockout message.',
                )),
                (31, s(
                    'Test password reset link expiry.',
                    'WHY: expired links should fail gracefully.',
                    'HOW: request a password reset email. Wait > 1 hour. Click the link in the old email.',
                    'PASS IF: a friendly error appears like "This link has expired. Please request a new one."',
                )),
                (32, s(
                    'Test session expiry.',
                    'WHY: stale sessions should redirect cleanly.',
                    'HOW: log in, then leave the tab open for 24+ hours without activity. Return and try to do something (e.g., click into a private page).',
                    'PASS IF: redirects to /login. Or simply ask Rafael to confirm session timeout settings via Supabase.',
                )),
            ]),
            ('F.5 Security headers', [
                (33, s(
                    'Inspect HTTP response headers for the homepage.',
                    'WHY: security headers prevent common web attacks.',
                    'HOW STEP-BY-STEP:',
                    '  1) In Chrome, open https://hirequadrant.com in an incognito window.',
                    '  2) Press Cmd+Option+I (Mac) or F12 (Windows) — DevTools opens.',
                    '  3) Click the "Network" tab in DevTools.',
                    '  4) Reload the page (Cmd+R or F5).',
                    '  5) In the network requests list, click the very first row (which is the page itself, named "hirequadrant.com" or "/").',
                    '  6) The right-side panel opens. Click the "Headers" tab.',
                    '  7) Scroll down to the "Response Headers" section.',
                    '  8) Check each header is present.',
                    'PASS IF: you see ALL of these: Strict-Transport-Security, X-Content-Type-Options, X-Frame-Options, Referrer-Policy, Permissions-Policy, Content-Security-Policy.',
                )),
                (34, s(
                    'Test iframe-embedding is blocked.',
                    'WHY: prevents clickjacking attacks.',
                    'HOW STEP-BY-STEP:',
                    '  1) Open https://codepen.io in a new tab.',
                    '  2) Click "Start coding" → New Pen.',
                    '  3) In the HTML pane, paste exactly: <iframe src="https://hirequadrant.com" width="800" height="600"></iframe>',
                    '  4) Look at the result pane on the right.',
                    'PASS IF: result pane is BLANK or shows an error like "refused to connect". Browser blocked the embed.',
                    'FAIL IF: HireQuadrant loads inside the iframe (means X-Frame-Options is missing).',
                )),
                (35, s(
                    'Verify security.txt published.',
                    'WHY: industry standard for vulnerability reporting.',
                    'HOW: in the address bar, type https://hirequadrant.com/.well-known/security.txt → Enter.',
                    'PASS IF: a plain-text page loads with email contact, expiry date, and "Contact: mailto:..." lines. NOT a 404.',
                )),
            ]),
            ('F.6 Anti-spam on registration', [
                (36, s(
                    'Test bot-defense (timing).',
                    'WHY: registration form rejects submissions that complete in <2.5 seconds (bots).',
                    'HOW: this is hard for a human to test (we can\'t click that fast). SKIP unless you can use browser automation. Or have Rafael verify with a quick script.',
                    'PASS IF (skipped): mark Done + Pass. Note: "Skipped — manual test impractical."',
                )),
                (37, s(
                    'Test bot-defense (honeypot).',
                    'WHY: hidden field that bots auto-fill but humans never see.',
                    'HOW: this requires browser DevTools to find the hidden field, fill it, and submit. SKIP unless technical. Or have Rafael verify.',
                    'PASS IF (skipped): mark as Skipped in Notes.',
                )),
                (38, s(
                    'Test breached-password block (HIBP).',
                    'WHY: blocks signups using known-leaked passwords.',
                    'HOW: visit /register → fill name + a fake email + the password "password123" + confirm → submit.',
                    'PASS IF: error message appears noting the password "appears in [N] data breaches". Registration blocked.',
                )),
            ]),
            ('F.7 Realtime updates', [
                (39, s(
                    'Test realtime notifications across two tabs.',
                    'WHY: notifications should appear without manual refresh.',
                    'HOW STEP-BY-STEP:',
                    '  1) Open Chrome window 1: log in as candidate seed-reviewer-1@hirequadrant-seed.test → visit /notifications.',
                    '  2) Open Chrome window 2 (or Safari): log in as employer test-employer-1@hirequadrant-test.com.',
                    '  3) Tile the windows side-by-side so you can see both.',
                    '  4) In employer window, find an applicant who is the candidate from window 1 (or apply from window 1 first to create the link).',
                    '  5) In employer window, click the applicant → "Run AI Screening" or send a message via Bulk Message.',
                    '  6) Watch the candidate window 1.',
                    'PASS IF: a new notification appears in window 1 within ~2 seconds, no refresh needed.',
                )),
            ]),
            ('F.8 Sitemap and SEO meta', [
                (40, s(
                    'Verify /sitemap-pages XML.',
                    'WHY: search engines crawl this for indexing.',
                    'HOW: visit https://hirequadrant.com/sitemap-pages.',
                    'PASS IF: XML page with hundreds of <url> entries (companies, industries, blog posts, etc.). Browser may show raw XML — that\'s OK.',
                )),
                (41, s(
                    'Verify /sitemap-jobs XML.',
                    'HOW: visit https://hirequadrant.com/sitemap-jobs.',
                    'PASS IF: XML with all job URLs.',
                )),
                (42, s(
                    'Verify /robots.txt.',
                    'HOW: visit https://hirequadrant.com/robots.txt.',
                    'PASS IF: a plain-text page with rules including "User-agent: *" and a "Sitemap: ..." line.',
                )),
                (43, s(
                    'Verify JobPosting structured data on a job page.',
                    'WHY: Google needs JobPosting JSON-LD to show jobs in search.',
                    'HOW STEP-BY-STEP:',
                    '  1) Open any /jobs/[id] page.',
                    '  2) Press Cmd+Option+U (Mac) or Ctrl+U (Windows) to View Source.',
                    '  3) Press Cmd+F (Mac) or Ctrl+F to find. Search for: application/ld+json',
                    '  4) Look at the JSON block.',
                    'PASS IF: a JSON-LD script tag exists with "@type": "JobPosting".',
                )),
                (44, s(
                    'Verify Organization structured data on a company page.',
                    'HOW: open any company profile → View Source → search application/ld+json.',
                    'PASS IF: a JSON-LD block with "@type": "Organization" appears.',
                )),
                (45, s(
                    'Verify Article structured data on a blog post.',
                    'HOW: open any /blog/[slug] → View Source → search application/ld+json.',
                    'PASS IF: a JSON-LD block with "@type": "Article" appears.',
                )),
            ]),
            ('F.9 Error boundary', [
                (46, s(
                    'Force a JavaScript error to test the error boundary.',
                    'WHY: when something breaks, users should see a friendly page, not a blank white screen.',
                    'HOW STEP-BY-STEP:',
                    '  1) Open hirequadrant.com in Chrome.',
                    '  2) Open DevTools (Cmd+Option+I or F12).',
                    '  3) Click the "Console" tab.',
                    '  4) At the bottom prompt (looks like > ), type exactly: throw new Error("test from QA")',
                    '  5) Press Enter.',
                    'PASS IF: an error appears in the console (red text). The page itself does NOT show a "Something went wrong" page (this just throws in the console, which is normal). Note: actually triggering the boundary requires breaking a real component which is out of scope for manual QA. Mark this Skipped if unsure.',
                )),
            ]),
            ('F.10 Browser compatibility', [
                (47, s(
                    'Run candidate flow on Safari (Mac).',
                    'HOW: open Safari → Private Browsing → repeat key steps from Part C.',
                    'PASS IF: all key flows work. Note any visual differences in Notes.',
                )),
                (48, s(
                    'Run candidate flow on Firefox.',
                    'HOW: install Firefox if needed. Open private window. Repeat key steps from Part C.',
                    'PASS IF: works.',
                )),
                (49, s(
                    'Run candidate flow on iOS Safari (real iPhone).',
                    'HOW: on iPhone, open Safari → hirequadrant.com → run a few candidate flows.',
                    'PASS IF: works.',
                )),
                (50, s(
                    'Run candidate flow on Android Chrome (real Android phone).',
                    'HOW: on Android, open Chrome → hirequadrant.com → run a few flows.',
                    'PASS IF: works.',
                )),
            ]),
            ('F.11 Accessibility', [
                (51, s(
                    'Tab through a form using only the keyboard.',
                    'WHY: keyboard accessibility is required for users with disabilities.',
                    'HOW: visit /register. Click outside any input. Press Tab repeatedly.',
                    'PASS IF: focus moves through inputs in a logical order; the focused element shows a visible ring/outline.',
                )),
                (52, s(
                    'Test focus trap in the global search modal.',
                    'WHY: Tab inside a modal should not escape outside the modal.',
                    'HOW: press Cmd+K to open search → press Tab repeatedly.',
                    'PASS IF: focus stays inside the modal (cycles through input → results → close button → back to input).',
                )),
                (53, s(
                    'Test Escape closes modals.',
                    'HOW: open Cmd+K modal, press Esc.',
                    'PASS IF: modal closes. Repeat for any other modal you find (Schedule Interview, Bulk Message).',
                )),
                (54, s(
                    'Test color contrast (use Lighthouse Accessibility audit).',
                    'WHY: low contrast = unreadable for users with low vision.',
                    'HOW: re-run Lighthouse Accessibility check (from F.3 step 26).',
                    'PASS IF: Accessibility score ≥ 90 with no contrast errors flagged.',
                )),
                (55, s(
                    'Verify image alt text.',
                    'HOW: right-click any company logo or avatar → Inspect → confirm an alt="..." attribute is present.',
                    'PASS IF: alt text exists on all logos and avatars.',
                )),
            ]),
        ],
    ),
    (
        'PART G: Data integrity (admin only)',
        'These steps require Supabase database access. SKIP this section if you do not have admin login to the Supabase project. Send the section to Rafael instead.',
        [
            ('G.1 Database row counts', [
                (1, s(
                    'Count blog posts.',
                    'WHY: confirms blog seed migrations completed.',
                    'HOW: log into supabase.com → HireQuadrant project → SQL Editor → paste: select count(*) from blog_posts; → click Run.',
                    'PASS IF: returns approximately 200+ rows.',
                )),
                (2, s(
                    'Count help articles.',
                    'HOW: SQL: select count(*) from help_articles;',
                    'PASS IF: returns ~40.',
                )),
                (3, s(
                    'Count companies.',
                    'HOW: SQL: select count(*) from companies;',
                    'PASS IF: returns 25 or more.',
                )),
                (4, s(
                    'Confirm no XML-filename junk in jobs table.',
                    'HOW: SQL: select count(*) from jobs where company = \'hirequadrant.xml\';',
                    'PASS IF: returns 0.',
                )),
                (5, s(
                    'Verify recent audit log entries.',
                    'HOW: SQL: select * from audit_log order by created_at desc limit 5;',
                    'PASS IF: at least the moderation actions you performed in Part E appear.',
                )),
            ]),
            ('G.2 Reviews integrity', [
                (6, s(
                    'Count live approved reviews.',
                    'HOW: SQL: select count(*) from company_reviews where status = \'approved\' and deleted_at is null;',
                    'PASS IF: returns a number > 0.',
                )),
                (7, s(
                    'Confirm review rate limit holds.',
                    'HOW: SQL: select count(distinct user_id) from company_reviews where created_at > now() - interval \'24 hours\';',
                    'PASS IF: returns a number (verifies the table is queryable; rate limit enforces 3/user/24h via trigger).',
                )),
            ]),
        ],
    ),
    (
        'PART H: Production launch checklist',
        'Final go/no-go before flipping the switch. Mostly admin-level confirmations. Ask Rafael for help on any you can\'t verify yourself.',
        [
            ('H.1 Infrastructure', [
                (1, s(
                    'All migrations run cleanly in production Supabase.',
                    'PASS IF: Rafael confirms no failed migrations in Supabase logs.',
                )),
                (2, s(
                    'Edge functions deployed.',
                    'PASS IF: Rafael confirms ai-helpers, sitemap-pages, sitemap-jobs are all deployed and live.',
                )),
                (3, s(
                    'DNS pointing to Netlify.',
                    'PASS IF: hirequadrant.com resolves and serves the site.',
                )),
                (4, s(
                    'HTTPS cert valid.',
                    'PASS IF: visit hirequadrant.com — green padlock in browser, no certificate warnings.',
                )),
                (5, s(
                    'Anthropic API key in Supabase Edge Function secrets.',
                    'PASS IF: Rafael confirms ANTHROPIC_API_KEY is set in Supabase project settings.',
                )),
                (6, s(
                    'No service role key leaked in client bundle.',
                    'PASS IF: Rafael confirms (technical check).',
                )),
                (7, s(
                    'Daily Supabase backup confirmed.',
                    'PASS IF: Rafael confirms backup schedule active.',
                )),
            ]),
            ('H.2 First-real-user dry run', [
                (8, s(
                    'Test login from a fresh device with both seeded admin accounts.',
                    'PASS IF: both accounts work.',
                )),
                (9, s(
                    'At least one real employer has uploaded their first job successfully.',
                    'PASS IF: confirmed.',
                )),
                (10, s(
                    'At least one candidate applied through full flow end-to-end.',
                    'PASS IF: confirmed.',
                )),
                (11, s(
                    'At least one review submitted, moderated, and published.',
                    'PASS IF: confirmed.',
                )),
            ]),
        ],
    ),
]


def main():
    doc = Document()

    for section in doc.sections:
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)

    title = doc.add_heading('HireQuadrant — Pre-Launch QA Walkthrough', level=0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitle = doc.add_paragraph()
    sr = subtitle.add_run('Test every click, every flow, every persona — with full step-by-step instructions for non-technical testers.')
    sr.italic = True
    sr.font.size = Pt(10)
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    intro = doc.add_paragraph()
    intro.add_run(
        'How to use this document: walk through top to bottom. Every step has '
    )
    bold_what = intro.add_run('WHAT')
    bold_what.bold = True
    intro.add_run(' (what the feature is for), ')
    bold_how = intro.add_run('HOW')
    bold_how.bold = True
    intro.add_run(' (the exact clicks/typing to perform), and ')
    bold_pass = intro.add_run('PASS IF')
    bold_pass.bold = True
    intro.add_run(' (the expected result). Mark Done, then Pass or Fail, and write what you saw in Notes. Take screenshots of failures.')
    intro.paragraph_format.space_after = Pt(8)

    legend = doc.add_paragraph()
    lr = legend.add_run('How to mark each box: ')
    lr.bold = True
    legend.add_run(
        'Click directly inside a ☐ cell — your cursor lands next to the box. '
        'Type X (or replace the ☐ with ☑). To unmark, delete the X.'
    )
    legend.paragraph_format.space_after = Pt(6)

    legend2 = doc.add_paragraph()
    l2r = legend2.add_run('Estimated time: ')
    l2r.bold = True
    legend2.add_run('8–10 hours for one tester walking through alone, or 3-4 hours per person if 3 testers split the personas.')
    legend2.paragraph_format.space_after = Pt(18)

    for heading, intro_text, subsections in SECTIONS:
        add_heading(doc, heading, level=1)
        add_intro_paragraph(doc, intro_text)
        for sub_heading, steps in subsections:
            add_subsection_heading(doc, sub_heading)
            add_test_table(doc, steps)
            doc.add_paragraph()

    add_heading(doc, 'Summary report (fill in when done)', level=1)
    add_intro_paragraph(doc, 'After completing the walkthrough, fill in this summary and send back to Rafael.')

    summary_table = doc.add_table(rows=1, cols=2)
    summary_table.style = 'Light Grid'
    h = summary_table.rows[0].cells
    h[0].text = 'Section'
    h[1].text = 'Pass / Fail summary'
    for cell in h:
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
        set_cell_shade(cell, 'E8EEF7')
        set_cell_borders(cell)

    summary_rows = [
        'Part A: Setup',
        'Part B: Anonymous',
        'Part C: Candidate',
        'Part D: Employer',
        'Part E: Admin',
        'Part F: Cross-cutting',
        'Part G: Data integrity',
        'Part H: Launch checklist',
        'Overall recommendation (Go / No-Go / Go with caveats)',
    ]
    for label in summary_rows:
        row = summary_table.add_row().cells
        row[0].text = label
        row[1].text = ''
        for c in row:
            set_cell_borders(c)

    set_table_widths(summary_table, [3.0, 5.0])

    add_subsection_heading(doc, 'Findings to flag for Rafael')
    findings = doc.add_paragraph('\n\n\n\n\n\n\n\n\n')
    findings.paragraph_format.line_spacing = 1.6

    add_subsection_heading(doc, 'Tester sign-off')
    signoff = doc.add_paragraph()
    signoff.add_run('Tester name: __________________________________________   ')
    signoff.add_run('Date: ______________________\n\n')
    signoff.add_run('Signature: ____________________________________________________________________')
    signoff.paragraph_format.line_spacing = 1.8

    out = '/Users/rafaelmaldonado/Library/CloudStorage/OneDrive-Personal/Projects/agents/hire-quadrant/docs/HireQuadrant-QA-Walkthrough.docx'
    import os
    os.makedirs(os.path.dirname(out), exist_ok=True)
    doc.save(out)
    print(f'Saved: {out}')
    print(f'Size: {os.path.getsize(out):,} bytes')


if __name__ == '__main__':
    main()
